"""
Convert the Tech Assessment markdown to a Word document using python-docx.
"""
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def parse_md_to_docx(md_path, docx_path):
    with open(md_path, 'r') as f:
        lines = f.readlines()

    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    i = 0
    in_table = False
    table_rows = []
    
    while i < len(lines):
        line = lines[i].rstrip('\n')
        
        # Skip empty lines
        if not line.strip():
            if in_table:
                # End of table, flush it
                _flush_table(doc, table_rows)
                in_table = False
                table_rows = []
            i += 1
            continue
        
        # Horizontal rule
        if line.strip() == '---':
            if in_table:
                _flush_table(doc, table_rows)
                in_table = False
                table_rows = []
            i += 1
            continue
        
        # Headers
        if line.startswith('#'):
            if in_table:
                _flush_table(doc, table_rows)
                in_table = False
                table_rows = []
            level = len(line) - len(line.lstrip('#'))
            text = line.lstrip('#').strip()
            heading = doc.add_heading(text, level=min(level, 4))
            i += 1
            continue
        
        # Table rows
        if '|' in line and line.strip().startswith('|'):
            stripped = line.strip()
            # Check if separator row
            if re.match(r'^\|[\s\-:|]+\|$', stripped):
                i += 1
                continue
            # Parse table row
            cells = [c.strip() for c in stripped.split('|')[1:-1]]
            if not in_table:
                in_table = True
                table_rows = []
            table_rows.append(cells)
            i += 1
            continue
        
        # If we were in a table and hit non-table content
        if in_table:
            _flush_table(doc, table_rows)
            in_table = False
            table_rows = []
        
        # Bullet points
        if line.strip().startswith('- ') or line.strip().startswith('* '):
            indent_level = (len(line) - len(line.lstrip())) // 2
            text = line.strip()[2:]
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.25 * (indent_level + 1))
            _add_formatted_text(p, text)
            i += 1
            continue
        
        # Numbered list
        if re.match(r'^\s*\d+\.\s', line):
            text = re.sub(r'^\s*\d+\.\s', '', line)
            p = doc.add_paragraph(style='List Number')
            _add_formatted_text(p, text)
            i += 1
            continue
        
        # Regular paragraph
        p = doc.add_paragraph()
        _add_formatted_text(p, line)
        i += 1
    
    # Flush any remaining table
    if in_table:
        _flush_table(doc, table_rows)
    
    doc.save(docx_path)
    print(f"Generated: {docx_path}")


def _flush_table(doc, table_rows):
    if not table_rows:
        return
    num_cols = max(len(row) for row in table_rows)
    table = doc.add_table(rows=len(table_rows), cols=num_cols)
    table.style = 'Table Grid'
    
    for r_idx, row in enumerate(table_rows):
        for c_idx, cell_text in enumerate(row):
            if c_idx < num_cols:
                cell = table.cell(r_idx, c_idx)
                cell.text = ''
                p = cell.paragraphs[0]
                _add_formatted_text(p, cell_text)
                # Bold the header row
                if r_idx == 0:
                    for run in p.runs:
                        run.bold = True


def _add_formatted_text(paragraph, text):
    """Parse markdown inline formatting and add runs to paragraph."""
    # Pattern for bold, italic, links, and inline code
    pattern = r'(\*\*.*?\*\*|\*.*?\*|_.*?_|`.*?`|\[.*?\]\(.*?\))'
    parts = re.split(pattern, text)
    
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif (part.startswith('*') and part.endswith('*')) or (part.startswith('_') and part.endswith('_')):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
        elif part.startswith('[') and '](' in part:
            # Link: [text](url)
            link_text = part[1:part.index(']')]
            run = paragraph.add_run(link_text)
            run.underline = True
        else:
            paragraph.add_run(part)


if __name__ == '__main__':
    parse_md_to_docx(
        'Documents/Shail_Bhatt_Tech_Assessment_L7_PTPM.md',
        'Documents/Shail_Bhatt_Tech_Assessment_L7_PTPM.docx'
    )
