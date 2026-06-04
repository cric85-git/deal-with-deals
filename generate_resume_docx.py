#!/usr/bin/env python3
"""Generate a formatted Word document from the revised resume."""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# Set narrow margins
sections = doc.sections
for section in sections:
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

# Helper functions
def set_paragraph_spacing(paragraph, before=0, after=0, line_spacing=1.15):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line_spacing

def add_horizontal_rule(doc, color="C4A44D"):
    """Add a colored horizontal rule."""
    paragraph = doc.add_paragraph()
    set_paragraph_spacing(paragraph, before=2, after=2)
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return paragraph

def add_section_heading(doc, text):
    """Add a styled section heading with gold underline."""
    heading = doc.add_paragraph()
    set_paragraph_spacing(heading, before=12, after=2)
    run = heading.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    add_horizontal_rule(doc)

def add_bullet(doc, bold_text, normal_text, indent=0.25):
    """Add a bullet point with bold lead and normal continuation."""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(indent)
    set_paragraph_spacing(p, before=4, after=4, line_spacing=1.15)
    if bold_text:
        run = p.add_run(bold_text)
        run.bold = True
        run.font.size = Pt(10)
    if normal_text:
        run = p.add_run(normal_text)
        run.font.size = Pt(10)
    return p

# ============ HEADER ============
name = doc.add_paragraph()
name.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = name.add_run("Shail Bhatt")
run.bold = True
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
set_paragraph_spacing(name, before=0, after=2)

contact = doc.add_paragraph()
contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = contact.add_run("1525 Alamosa Dr, Plano, TX 75023 | (832) 622-4681 | shailbhatt@yahoo.com")
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
set_paragraph_spacing(contact, before=0, after=1)

linkedin = doc.add_paragraph()
linkedin.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = linkedin.add_run("LinkedIn: www.linkedin.com/in/shailbhatt")
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
set_paragraph_spacing(linkedin, before=0, after=6)

add_horizontal_rule(doc)

# ============ PROFESSIONAL SUMMARY ============
add_section_heading(doc, "Professional Summary")

p = doc.add_paragraph()
set_paragraph_spacing(p, before=4, after=4, line_spacing=1.15)
run = p.add_run(
    "Technology executive who builds and scales AI-first platforms that transform how organizations "
    "detect, diagnose, and resolve problems at scale. 15+ years leading enterprise programs at Amazon "
    "and Verizon — from vision through execution — delivering $100M+ in measurable business impact "
    "across observability, customer journey analytics, GenAI platforms, and cloud infrastructure."
)
run.font.size = Pt(10)

p = doc.add_paragraph()
set_paragraph_spacing(p, before=4, after=4, line_spacing=1.15)
run = p.add_run(
    "I bring together engineering, product, data science, and executive leadership around shared "
    "outcomes — turning ambiguous, cross-organizational challenges into clear roadmaps, operating "
    "models, and shipped results. Two US patents. IEEE Senior Member."
)
run.font.size = Pt(10)

# ============ LEADERSHIP SCOPE ============
add_section_heading(doc, "Leadership Scope")

scope_items = [
    ("Organizational Influence: ", "5+ VP-level organizations across Device OS, Services, Product, and Data Science"),
    ("Program Scale: ", "Multi-year, enterprise-wide transformation programs spanning 50+ teams and 3 major device product lines"),
    ("Executive Engagement: ", "Regular strategy, investment, and roadmap presentations to VP/SVP leadership (including Product Day for SVP Panos Panay)"),
    ("Team Building: ", "Built and mentored global cross-functional teams across engineering, PM, UX, and data science — operating as both TPM and SDM when needed"),
    ("Investment Influence: ", "Shaped GenAI and Agentic AI investment priorities and multi-year roadmaps with VP leadership"),
]

for bold, normal in scope_items:
    add_bullet(doc, bold, normal)

# ============ WORK HISTORY ============
add_section_heading(doc, "Work History")

# Amazon
role = doc.add_paragraph()
set_paragraph_spacing(role, before=8, after=1)
run = role.add_run("Senior Technical Program Manager — Amazon")
run.bold = True
run.font.size = Pt(11)

location = doc.add_paragraph()
set_paragraph_spacing(location, before=0, after=6)
run = location.add_run("Dallas, TX | 2021 – Present")
run.font.size = Pt(10)
run.italic = True
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

amazon_bullets = [
    ("Conceived and delivered a unified AI-driven troubleshooting platform ",
     "spanning detection, diagnosis, and remediation across FireTV, Tablets, and Alexa — reducing MTTR by 30%, improving first-attempt resolution by 55%, and achieving 500+ organic customers within 3 weeks of launch. Took over a greenfield DSLT goal mid-flight with no SDM, wore both TPM and SDM hats, and delivered 13 days ahead of deadline."),
    ("Authored a PRFAQ to solve an ambiguous, cross-domain diagnosis problem ",
     "— proposing an integrated system that automatically correlates telemetry across device logs, crashes, user interactions, cloud service logs, customer support data, and system traces into a unified diagnostic ledger. Drove the vision through 3 months of stakeholder interviews, PE reviews, and L8/L10 sign-off — then translated it into a production Diagnosis MCP (Model Context Protocol) tool that puts end-to-end automated root-cause analysis into action across multiple domains."),
    ("Owned enterprise-wide telemetry platform consolidation ",
     "across device OS, services, and product organizations — driving migration of 98% of 11.2T+ daily metrics to a unified, privacy-first observability architecture, enabling AI-driven reliability, compliance, and large-scale experimentation."),
    ("Launched a customer journey and lifecycle analytics platform ",
     "providing end-to-end visibility into device usage and engagement — informing product strategy, hardware investment, and experience optimization across 3 major product lines."),
    ("Drove large-scale cloud modernization, ",
     "leading migration of legacy systems to AWS and delivering $22.4M in annualized infrastructure savings while improving scalability and operational resilience."),
    ("Partnered with VP leadership to define GenAI and Agentic AI strategy, ",
     "multi-year roadmaps, and investment priorities — translating emerging AI capabilities into production platforms and business outcomes, including a live agentic troubleshooting demo for SVP and 20+ L8–L10 leaders."),
    ("Established cross-organization launch readiness and execution mechanisms, ",
     "defining planning rhythms, risk management, and release governance that accelerated high-quality launches across device, application, and service portfolios."),
    ("Architected a cross-platform middleware strategy ",
     "based on a shared native-code framework to support a \"one bug, one fix\" model — reducing engineering duplication and accelerating release velocity across all device platforms."),
]

for bold, normal in amazon_bullets:
    add_bullet(doc, bold, normal)

# Verizon
role = doc.add_paragraph()
set_paragraph_spacing(role, before=12, after=1)
run = role.add_run("Senior Manager, System of Insights — Verizon")
run.bold = True
run.font.size = Pt(11)

location = doc.add_paragraph()
set_paragraph_spacing(location, before=0, after=6)
run = location.add_run("Irving, TX | 2010 – 2021")
run.font.size = Pt(10)
run.italic = True
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

verizon_bullets = [
    ("Delivered $69M+ in business impact ",
     "through data-driven initiatives — including 12M call reductions, 79.7% digital self-serve rate, and targeted engagement across the Wireless segment."),
    ("Designed and deployed a real-time data platform ",
     "processing 150M+ daily events — enabling unified customer journey analytics, omni-channel orchestration, and proposition arbitration."),
    ("Led AI/ML strategy delivering patented fraud detection and predictive billing models ",
     "— boosting agent efficiency, increasing digital deflection by 45%, and reducing call-in rates by 27%."),
    ("Mentored and led high-performing global teams, ",
     "driving customer journey strategies, personalized experiences, omni-channel integration, and predictive Next-Best Actions at enterprise scale."),
]

for bold, normal in verizon_bullets:
    add_bullet(doc, bold, normal)

# ============ INNOVATION & RECOGNITION ============
add_section_heading(doc, "Innovation & Recognition")

innovation_items = [
    ("US Patent: ", "Systems and methods for identifying a service qualification of a unit based on an image-based window analysis. Patent #11521369 (December 2022)"),
    ("US Patent: ", "Chat Analysis Using Machine Learning. Patent #10990762 (April 2021)"),
    ("IEEE Senior Member ", "— Elevated in 2025"),
    ("Winner, Startup Weekend Denton 2014 ", "— Pitched and prototyped a food discovery app that led to the launch of FiveOh, Inc."),
]

for bold, normal in innovation_items:
    add_bullet(doc, bold, normal)

# ============ EDUCATION ============
add_section_heading(doc, "Education")

edu_items = [
    ("University of Texas at Dallas ", "— M.S. Computer Science (2009)"),
    ("Charotar Institute of Science and Technology ", "— B.E. Information Technology (2007)"),
]

for bold, normal in edu_items:
    add_bullet(doc, bold, normal)

# ============ AREAS OF IMPACT ============
add_section_heading(doc, "Areas of Impact")

areas = [
    "AI/ML Platform Strategy & Productionization (GenAI, Agentic AI, LLMs, MCP)",
    "Enterprise Observability & Reliability at Scale",
    "Cross-Organizational Program Leadership & Stakeholder Alignment",
    "Cloud Modernization & Infrastructure Economics",
    "Customer Journey Analytics & Experience Optimization",
    "Ambiguity Navigation — Defining Problems, Building Consensus, Shipping Solutions",
]

for area in areas:
    add_bullet(doc, "", area)

# Save
output_path = "Documents/Shail_Bhatt_Resume_Revised.docx"
doc.save(output_path)
print(f"✅ Word document saved to: {output_path}")
